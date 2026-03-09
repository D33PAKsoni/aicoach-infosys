import os
import io
import uuid
import json
import re
import numpy as np
import whisper
import pypdf
import docx
import anyio 
from fastapi import APIRouter, File, UploadFile, Form, HTTPException, Depends, Request
from fastapi.responses import FileResponse
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from pydub import AudioSegment
import speech_recognition as sr
from gtts import gTTS
import pyttsx3
from google import genai
from google.genai import types
from .. import models
from ..core.config import settings
from sqlalchemy.orm import Session
from ..database import SessionLocal

router = APIRouter()

nlp_model = SentenceTransformer('all-mpnet-base-v2', device="cpu")
whisper_model = whisper.load_model("base")
client = genai.Client(api_key=settings.GEMINI_API_KEY)

MODEL_ID = "gemini-2.5-flash" 
chat_sessions = {}
AUDIO_DIR = "temp_audio"
os.makedirs(AUDIO_DIR, exist_ok=True)


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def get_embedding(text):
    words = text.split()
    chunk_size = 300
    chunks = [' '.join(words[i:i + chunk_size]) for i in range(0, len(words), chunk_size)]
    if not chunks: return np.zeros(768)
    chunk_embeddings = nlp_model.encode(chunks)
    return np.mean(chunk_embeddings, axis=0)

async def extract_text_from_file(file: UploadFile):
    content = await file.read()
    if file.filename.endswith(".pdf"):
        pdf_reader = pypdf.PdfReader(io.BytesIO(content))
        return "".join([page.extract_text() or "" for page in pdf_reader.pages])
    elif file.filename.endswith(".docx"):
        doc = docx.Document(io.BytesIO(content))
        return "\n".join([para.text for para in doc.paragraphs])
    elif file.filename.endswith(".txt"):
        return content.decode("utf-8")
    return ""


def extract_text_from_path(file_path: str):
    if not os.path.exists(file_path):
        return ""
    
    if file_path.endswith(".pdf"):
        with open(file_path, "rb") as f:
            pdf_reader = pypdf.PdfReader(f)
            return "".join([page.extract_text() or "" for page in pdf_reader.pages])
    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    return ""

def extract_json(text):
    match = re.search(r'\{.*\}', text, re.DOTALL)
    return match.group(0) if match else text

@router.post("/analyze")
async def analyze_match(resume: UploadFile = File(...), jd_file: UploadFile = File(None), jd_text: str = Form(None)):
    if not jd_file and not jd_text:
        raise HTTPException(status_code=400, detail="Provide JD file or text.")
    
    resume_content = await extract_text_from_file(resume)
    jd_content = await extract_text_from_file(jd_file) if jd_file else jd_text

    if not resume_content.strip() or not jd_content.strip():
        raise HTTPException(status_code=400, detail="Extraction failed.")

    r_emb = await anyio.to_thread.run_sync(get_embedding, resume_content)
    j_emb = await anyio.to_thread.run_sync(get_embedding, jd_content)
    
    score = cosine_similarity([r_emb], [j_emb])[0][0]
    
    return {"match_percentage": round(float(score) * 100, 2), "status": "Success"}



@router.post("/aianalyze")
async def analyze_match(resume: UploadFile = File(...), jd_file: UploadFile = File(None), jd_text: str = Form(None)):
    if not jd_file and not jd_text:
        raise HTTPException(status_code=400, detail="Provide JD file or text.")
    
    resume_content = await extract_text_from_file(resume)
    jd_content = await extract_text_from_file(jd_file) if jd_file else jd_text

    if not resume_content.strip() or not jd_content.strip():
        raise HTTPException(status_code=400, detail="Extraction failed.")
    
    prompt = (
        f"You are an expert HR recruiter. Analyze the match between the following Resume and Job Description (JD).\n\n"
        f"### Resume:\n{resume_content}\n\n"
        f"### Job Description:\n{jd_content}\n\n"
        "Return a JSON object with the following fields:\n"
        "- match_percentage (int): overall compatibility score from 0-100\n"
        "- summary (str): a 2-3 sentence overview of the candidate's fit\n"
        "- key_matches (list): specific skills or experiences that align with the JD\n"
        "- missing_skills (list): critical requirements from the JD not found in the resume\n"
        "- suggestions (list): tips to improve the resume for this specific role\n"
        "Provide ONLY the JSON object, no introductory text."
    )

    try:
        response = client.models.generate_content(
            model=MODEL_ID,
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json"
            )
        )

        analysis_data = json.loads(extract_json(response.text))
        return {"match_percentage": round(float(analysis_data["match_percentage"]),2), "status": "Success"}
        # return {**analysis_data, "status": "Success"}

    except Exception as e:
        # Fallback if the AI fails or JSON parsing errors out
        raise HTTPException(status_code=500, detail=f"AI Analysis failed: {str(e)}")




 

@router.get("/resumes")
async def get_user_resumes(user_id: int, db: Session = Depends(get_db)):
    return db.query(models.Resume).filter(models.Resume.user_id == user_id).all()

@router.get("/resumes/download/{resume_id}")
async def download_resume(resume_id: int, db: Session = Depends(get_db)):
    res_db = db.query(models.Resume).filter(models.Resume.id == resume_id).first()
    if not res_db or not os.path.exists(res_db.file_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(res_db.file_path, filename=res_db.file_name)

@router.post("/upload-resume")
async def upload_resume(user_id: int = Form(...), file: UploadFile = File(...), db: Session = Depends(get_db)):
    upload_dir = "uploads/resumes"
    os.makedirs(upload_dir, exist_ok=True)
    
    file_path = os.path.join(upload_dir, f"{uuid.uuid4()}_{file.filename}")
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())
    
    new_resume = models.Resume(user_id=user_id, file_name=file.filename, file_path=file_path)
    db.add(new_resume)
    db.commit()
    return {"message": "Resume saved to library"}

@router.delete("/resumes/{resume_id}")
async def delete_resume(resume_id: int, db: Session = Depends(get_db)):
    res_db = db.query(models.Resume).filter(models.Resume.id == resume_id).first()
    if not res_db:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    if os.path.exists(res_db.file_path):
        os.remove(res_db.file_path)
        
    db.delete(res_db)
    db.commit()
    return {"message": "Resume deleted successfully"}



@router.post("/transcribe")
def transcribe_audio(audio: UploadFile = File(...)):
    file_id = str(uuid.uuid4())
    temp_path = os.path.join(AUDIO_DIR, f"{file_id}_raw")
    pcm_path = os.path.join(AUDIO_DIR, f"{file_id}_fixed.wav")
    
    with open(temp_path, "wb") as f:
        f.write(audio.file.read())

    try:
        audio_segment = AudioSegment.from_file(temp_path)
        duration = audio_segment.duration_seconds
        audio_segment.export(pcm_path, format="wav")
        
        try:
            r = sr.Recognizer()
            with sr.AudioFile(pcm_path) as source:
                audio_data = r.record(source)
                transcript = r.recognize_google(audio_data)
        except:
            result = whisper_model.transcribe(pcm_path)
            transcript = result['text']

        return {"transcript": transcript, "duration": duration}
    except Exception as e:
        return {"transcript": "", "duration": 0}
    finally:
        for path in [temp_path, pcm_path]:
            if os.path.exists(path): os.remove(path)




@router.post("/speak")
async def speak_text(text: str = Form(...)):
    file_id = str(uuid.uuid4())
    audio_path = os.path.join(AUDIO_DIR, f"{file_id}.mp3")
    success = False

    try:
        engine = pyttsx3.init()
        wav_path = os.path.join(AUDIO_DIR, f"{file_id}.wav")
        engine.save_to_file(text, wav_path)
        engine.runAndWait()
        if os.path.exists(wav_path):
            audio_path, success = wav_path, True
    except: pass

    if not success:
        try:
            tts = gTTS(text=text, lang='en')
            tts.save(audio_path)
            success = True
        except:
            raise HTTPException(status_code=500, detail="TTS failed.")

    return FileResponse(audio_path, media_type="audio/mpeg" if audio_path.endswith(".mp3") else "audio/wav")


@router.post("/cleanup-audio")
async def cleanup_audio():
    try:
        for filename in os.listdir(AUDIO_DIR):
            file_path = os.path.join(AUDIO_DIR, filename)
            if os.path.isfile(file_path):
                os.remove(file_path)
        return {"status": "Temporary audio cleared"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Cleanup failed: {str(e)}")



@router.post("/generate-question")
def generate_question(transcript: str = Form(...), session_id: str = Form(...), resume_context: str = Form(""), jd_context: str = Form("")):
    try:
        if session_id not in chat_sessions:
            system_instruction = (
                f"You are a technical interviewer for: {jd_context}. Resume: {resume_context}. "
                "Return ONLY a JSON object: {'question': str, 'accuracy': int, 'filler_words': list}"
            )
            chat_sessions[session_id] = client.chats.create(
                model=MODEL_ID,
                config=types.GenerateContentConfig(system_instruction=system_instruction)
            )

        response = chat_sessions[session_id].send_message(transcript)
        analysis = json.loads(extract_json(response.text))
        return analysis
    except Exception as e:
        return {"question": "Could not generate question.", "accuracy": 0, "filler_words": []}



@router.post("/multiplematch")
async def match_multiple_resumes(
    user_id: int = Form(...), 
    jd_file: UploadFile = File(None), 
    jd_text: str = Form(None),
    db: Session = Depends(get_db)
):

    if not jd_file and not jd_text:
        raise HTTPException(status_code=400, detail="Provide JD file or text.")
    
    jd_content = await extract_text_from_file(jd_file) if jd_file else jd_text
    if not jd_content.strip():
        raise HTTPException(status_code=400, detail="JD content is empty.")

    j_emb = await anyio.to_thread.run_sync(get_embedding, jd_content)

    user_resumes = db.query(models.Resume).filter(models.Resume.user_id == user_id).all()
    
    if not user_resumes:
        return {"results": [], "message": "No resumes found for this user."}

    results = []
    
    for res in user_resumes:
        try:
            resume_text = await anyio.to_thread.run_sync(extract_text_from_path, res.file_path)
            
            if not resume_text.strip():
                results.append({
                    "resume_id": res.id,
                    "file_name": res.file_name,
                    "match_percentage": 0,
                    "error": "Could not extract text"
                })
                continue

            r_emb = await anyio.to_thread.run_sync(get_embedding, resume_text)
            score = cosine_similarity([r_emb], [j_emb])[0][0]
            
            results.append({
                "resume_id": res.id,
                "file_name": res.file_name,
                "match_percentage": round(float(score) * 100, 2)
            })
        except Exception as e:
            results.append({
                "resume_id": res.id,
                "file_name": res.file_name,
                "match_percentage": 0,
                "error": str(e)
            })

    return {"results": results, "status": "Success"}


@router.post("/aimultianalyse")
async def ai_multiple_match(
    user_id: int = Form(...), 
    jd_file: UploadFile = File(None), 
    jd_text: str = Form(None),
    db: Session = Depends(get_db)
):
    # 1. Extract JD Text
    if not jd_file and not jd_text:
        raise HTTPException(status_code=400, detail="Provide JD file or text.")
    
    jd_content = await extract_text_from_file(jd_file) if jd_file else jd_text
    if not jd_content.strip():
        raise HTTPException(status_code=400, detail="JD Extraction failed.")

    # 2. Retrieve all resumes for this user from DB
    user_resumes = db.query(models.Resume).filter(models.Resume.user_id == user_id).all()
    if not user_resumes:
        return {"results": [], "message": "No resumes found for this user."}

    # 3. Extract text from all resumes
    resume_data = []
    for res in user_resumes:
        text = await anyio.to_thread.run_sync(extract_text_from_path, res.file_path)
        if text.strip():
            resume_data.append({"id": res.id, "name": res.file_name, "content": text})

    if not resume_data:
        return {"results": [], "message": "No readable resumes found."}

    # 4. Construct Prompt for Gemini
    # We send all resume summaries in one go for comparison if the list is small, 
    # or iterate if the list is very large. Here we process them together for ranking.
    resumes_prompt = "\n\n".join([f"RESUME ID {r['id']} ({r['name']}):\n{r['content']}" for r in resume_data])
    
    prompt = (
        f"You are an AI Recruitment Assistant. Compare the following list of resumes against the Job Description (JD).\n\n"
        f"### Job Description:\n{jd_content}\n\n"
        f"### Candidate Resumes:\n{resumes_prompt}\n\n"
        "Return a JSON object with a key 'results' containing a list of objects. Each object must have:\n"
        "- resume_id (int): the ID provided above\n"
        "- file_name (str): the name of the file provided above\n"
        "- match_percentage (int): 0-100 score\n"
        "- brief_reason (str): why this candidate is or isn't a fit\n"
        "- top_skills (list): matching skills found\n"
        "- improvement_suggestions (str): suggestions for improving the resume\n"
        "Sort the list by match_percentage in descending order. Provide ONLY JSON."
    )

    try:
        # 5. Call Gemini API
        response = client.models.generate_content(
            model=MODEL_ID,
            contents=prompt,
            config=types.GenerateContentConfig(response_mime_type="application/json")
        )
        
        analysis_data = json.loads(extract_json(response.text))
        # return {**analysis_data, "status": "Success"}
        return {"results": analysis_data.get("results", []), "status": "Success"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI Multiple Analysis failed: {str(e)}")